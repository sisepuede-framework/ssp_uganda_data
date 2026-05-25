"""
generate_cba_report.py
Reads cba_formula_mapping.csv and writes cba_methodology_report.docx.
Structure: 5-section main body + 2 appendix tables.
Run in: ssp_uganda_cb
"""

import pandas as pd
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
CSV  = HERE / "cba_formula_mapping.csv"
OUT  = HERE / "cba_methodology_report.docx"

# ── Colours ────────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x39, 0x64)
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
GREY_FILL  = "D9E1F2"
LIGHT_FILL = "EEF3FB"

# ── Sector display order for Appendix A ───────────────────────────────────────
SECTOR_ORDER = [
    "AG - Crops", "AG - Livestock", "AG - Livestock (manure)", "CCS",
    "EN - Building", "EN - Electricity/Heat (enfu)", "EN - Fugitive emissions",
    "EN - Industrial combustion", "EN - Power Industry (entc)", "EN - Transportation",
    "IN - Industrial processes", "LULUCF - Forest land", "LULUCF - Organic soil",
    "Waste - Solid waste", "Waste - Wastewater (liquid waste)",
    "Waste - Wastewater treatment (trww)",
]

BENEFIT_ORDER = [
    "Air pollution (indoor)", "Consumer savings", "Crop value",
    "Ecosystem services", "Ecosystem services grasslands", "Ecosystem services wetlands",
    "Fuel cost savings", "Human Health", "IPPU value", "Livestock value",
    "O&M savings", "Other costs/benefits", "Pollution (air)",
    "Pollution (environment)", "Pollution (land)", "Pollution (water)",
]

# ── XML helpers ────────────────────────────────────────────────────────────────

def _set_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


# ── Document helpers ───────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE if level <= 2 else MID_BLUE
    run.font.bold = True
    run.font.size = Pt({1: 15, 2: 12, 3: 11}.get(level, 11))
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.size = Pt(10)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_formula_box(doc, lines):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_bg(cell, "F2F2F2")
    cell.paragraphs[0].clear()
    for line in lines:
        p = cell.add_paragraph(line)
        if p.runs:
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(9.5)
        else:
            run = p.add_run("")
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def add_inline_table(doc, headers, rows_data, col_widths):
    """Small inline table for examples in the body text."""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        _set_bg(hdr.cells[i], GREY_FILL)
        run = hdr.cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_BLUE
    for idx, row_vals in enumerate(rows_data):
        tr = tbl.add_row()
        fill = LIGHT_FILL if idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row_vals):
            _set_bg(tr.cells[i], fill)
            tr.cells[i].paragraphs[0].add_run(str(val)).font.size = Pt(9)
    _set_col_widths(tbl, col_widths)
    doc.add_paragraph()


def add_appendix_table(doc, df, col_defs, group_col=None):
    """
    col_defs: list of (csv_col, header, width_cm)
    If group_col is given, merge repeated values in that column visually (bold header rows).
    """
    headers = [h for _, h, _ in col_defs]
    widths  = [w for _, _, w in col_defs]
    cols    = [c for c, _, _ in col_defs]

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        _set_bg(hdr.cells[i], GREY_FILL)
        run = hdr.cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_BLUE

    prev_group = None
    for row_idx, (_, row) in enumerate(df.iterrows()):
        tr = tbl.add_row()
        group_val = row.get(group_col, "") if group_col else None
        is_new_group = group_col and group_val != prev_group

        fill = "E8EEF7" if is_new_group else (LIGHT_FILL if row_idx % 2 == 0 else "FFFFFF")
        for i, col in enumerate(cols):
            _set_bg(tr.cells[i], fill)
            val = "" if pd.isna(row.get(col, "")) else str(row.get(col, ""))
            # Suppress repeated group values to reduce visual noise
            if group_col and col == group_col and not is_new_group:
                val = ""
            run = tr.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            if is_new_group and i == 0:
                run.font.bold = True

        prev_group = group_val

    _set_col_widths(tbl, widths)
    doc.add_paragraph()


# ── Main ───────────────────────────────────────────────────────────────────────

def build():
    df = pd.read_csv(CSV)
    df["transformation_code"] = df["transformation_code"].fillna("")

    costs   = df[df["tableau_type"] == "Cost"].copy()
    benefits = df[df["tableau_type"] == "Benefit"].copy()

    # Forest Land example rows for Section 3
    forest_costs = costs[costs["tableau_category"] == "LULUCF - Forest land"].copy()
    forest_benefits = benefits[benefits["sector_code"] == "lndu"].copy()

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = title.add_run("Uganda Climate Cost-Benefit Analysis\n")
    r1.bold = True; r1.font.size = Pt(18); r1.font.color.rgb = DARK_BLUE
    r2 = title.add_run("Methodology and Factor Reference")
    r2.font.size = Pt(13); r2.font.color.rgb = MID_BLUE

    meta = doc.add_paragraph(f"Generated: {date.today().strftime('%B %d, %Y')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in meta.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()

    # ── Section 1: Overview ────────────────────────────────────────────────────
    add_heading(doc, "1. Overview", 1)
    add_body(doc, (
        "This document provides a complete methodological reference for the cost-benefit analysis (CBA) "
        "presented in Uganda's Climate Policy dashboard. The analysis translates outputs from SISEPUEDE — "
        "an integrated sectoral emissions modelling framework — into monetary values, covering the period "
        "2015 to 2070. All values are expressed in billions of PPP 2019 USD and normalised to GDP for "
        "cross-strategy comparability."
    ))
    add_body(doc, (
        "Each cost or benefit item in the dashboard is derived from one of two configuration sources:"
    ))
    add_bullet(doc, (
        "System-level factors (cost_factors): applied to any strategy that changes the relevant SISEPUEDE "
        "activity variable, regardless of which specific transformation caused the change. Examples include "
        "crop production value, ecosystem service values, fuel expenditure, and air pollution damages."
    ))
    add_bullet(doc, (
        "Investment costs (transformation_costs): applied only when a specific transformation is active "
        "in a strategy (e.g. TX:LNDU:INC_REFORESTATION). This configuration source contains two "
        "sub-types distinguished by the cb_type field in the variable name:"
    ))
    add_bullet(doc, (
        "      •  technical_cost (37 items): upfront capital and implementation expenditure required "
        "to deploy a transformation. Factors are typically negative (money spent). "
        "These appear under the Capital cost category in Tableau, broken down by sector."
    ))
    add_bullet(doc, (
        "      •  technical_savings (9 items): ongoing operational savings generated once a "
        "transformation is in place — for example, reduced maintenance costs from electric vehicles "
        "or heat pumps. These appear under the O&M savings category in Tableau as a co-benefit."
    ))
    add_body(doc, (
        "Both sources feed the same Tableau view and their values are summed within each category. "
        "The complete disaggregation — showing every line item, its source, unit cost factor, and "
        "formula — is provided in Appendices A and B."
    ))

    # ── Section 2: Computation Formulas ───────────────────────────────────────
    add_heading(doc, "2. Computation Formulas", 1)
    add_body(doc, (
        "Three formula types are used across all 167 cost and benefit line items. "
        "The formula type for each item is shown in the appendix tables."
    ))

    add_heading(doc, "2.1  Standard formula (101 of 167 items)", 2)
    add_body(doc, (
        "The large majority of items — including all crop and livestock production values, "
        "fuel cost savings, pollution damages, ecosystem services, and most capital costs — "
        "use the following formula:"
    ))
    add_formula_box(doc, [
        "value (USD) = (SSP_strategy_variable − SSP_baseline_variable)",
        "              × cost_factor",
        "              × annual_change ^ max(0, year − 2023)",
        "",
        "  SSP_strategy_variable   SISEPUEDE output for the policy scenario",
        "  SSP_baseline_variable   SISEPUEDE output for the BAU (business-as-usual) baseline",
        "  cost_factor             Monetary unit cost (e.g. USD/tonne, USD/capita, USD/PJ)",
        "  annual_change           Year-on-year scaling factor; 1.0 = constant over time",
        "  year                    Calendar year; scaling applied from 2023 onward",
    ])
    add_body(doc, (
        "Sign convention: the sign of the cost factor indicates the direction of the monetary flow. "
        "A negative factor produces a cost (capital investment, damage, or expenditure). "
        "A positive factor produces a gain (revenue, avoided damage, or economic benefit). "
        "The factor sign, combined with the direction of change in the SISEPUEDE variable, "
        "determines the final sign of each value in the output. Both costs and benefits "
        "can appear within the same Tableau category depending on the net effect of the strategy."
    ))

    add_heading(doc, "2.2  Comparison-strategy formula (29 items)", 2)
    add_body(doc, (
        "Used when a transformation's cost or benefit only makes economic sense relative to another "
        "policy scenario rather than the BAU baseline. This applies to industrial fuel switching, "
        "vehicle electrification, reforestation, and several other transformations that are evaluated "
        "against a renewable electricity scenario."
    ))
    add_formula_box(doc, [
        "value (USD) = (SSP_strategy_variable − SSP_comparison_strategy_variable)",
        "              × cost_factor",
        "              × annual_change ^ max(0, year − 2023)",
    ])
    add_body(doc, (
        "The comparison strategy for each item is specified in the configuration and is shown "
        "in the transformation code column of the appendix tables."
    ))

    add_heading(doc, "2.3  Sector-specific functions (14 items)", 2)
    add_body(doc, (
        "A subset of items use custom calculation logic in which the activity variable is "
        "derived differently from a direct SSP output difference. The unit cost factor shown "
        "in the appendix is still applied, but the quantity it multiplies is computed by "
        "sector-specific methods. Key cases:"
    ))
    add_bullet(doc, (
        "WALI Sanitation: costs and health benefits are applied to the population transitioning "
        "between sanitation tiers (unimproved → improved → safely managed). "
        "See Section 5.1 for the per-capita factors."
    ))
    add_bullet(doc, (
        "Livestock (enteric fermentation, manure management): reduction in emissions per "
        "tropical livestock unit, calibrated to Uganda livestock populations."
    ))
    add_bullet(doc, (
        "Fugitive emissions abatement: costs follow an abatement cost curve applied to "
        "changes in fugitive emission activity."
    ))
    add_bullet(doc, (
        "Industrial processes (IPPU clinker, fluorinated gases, CCS): sector model outputs "
        "for cement production, industrial gas substitution, and carbon capture."
    ))

    # ── Section 3: Stacking ────────────────────────────────────────────────────
    add_heading(doc, "3. How Costs and Benefits Are Aggregated", 1)
    add_body(doc, (
        "Within a given strategy, the same Tableau category can receive contributions from both "
        "system-level factors and investment costs. These are independent calculations that are "
        "concatenated in the pipeline and summed by Tableau when the dashboard is rendered."
    ))
    add_body(doc, (
        "The Forest Land sector illustrates this clearly. For a strategy that activates "
        "TX:LNDU:INC_REFORESTATION, four line items contribute:"
    ))

    # Hard-coded 4-row example: 2 system-level benefits + 2 investment costs
    forest_example_rows = [
        [
            "System cost",
            "Economic value of ecosystem services provided by retained primary forest "
            "(biodiversity, water regulation, carbon sequestration proxy)",
            "+500 $/ha",
            "Ecosystem services (Benefit)",
        ],
        [
            "System cost",
            "Economic value of ecosystem services recovered when pasture converts "
            "to secondary forest through reforestation",
            "+300 $/ha",
            "Ecosystem services (Benefit)",
        ],
        [
            "Investment cost\n(TX:LNDU:INC_SILVOPASTURE)",
            "Capital and implementation cost of establishing silvopasture "
            "(integrating trees into livestock grazing areas)",
            "−45 $/ha",
            "Capital cost (Cost)",
        ],
        [
            "Investment cost\n(TX:LNDU:INC_REFORESTATION)",
            "Capital and implementation cost of restoring degraded land "
            "through reforestation",
            "−88 $/ha",
            "Capital cost (Cost)",
        ],
    ]

    add_inline_table(doc,
        ["Source", "Description", "Factor", "Tableau category"],
        forest_example_rows,
        [3.5, 7.5, 1.8, 3.5],
    )

    add_body(doc, (
        "Tableau sums all Capital cost rows within the 'LULUCF — Forest land' sector to produce "
        "the cost bar, and sums all Ecosystem services rows separately to produce the benefit bar. "
        "Strategy interactions (where multiple transformations affect the same variable) are handled "
        "by a post-processing step that rescales values to avoid double-counting, using the weights "
        "defined in the strategy_interactions configuration sheet."
    ))

    # ── Section 4: GDP Normalization ──────────────────────────────────────────
    add_heading(doc, "4. GDP Normalisation", 1)
    add_body(doc, (
        "All monetary values produced by the pipeline are in billions of PPP 2019 USD. "
        "Normalisation to GDP is applied in Tableau:"
    ))
    add_formula_box(doc, [
        "share_of_GDP = value_B_USD / gdp_mmm_usd",
        "",
        "  value_B_USD   Computed cost or benefit, billions of PPP 2019 USD",
        "  gdp_mmm_usd   Uganda GDP, millions of PPP 2019 USD (annual, from model inputs)",
        "",
        "  Note: dividing billions by millions gives a dimensionless share,",
        "        displayed as a percentage in Tableau.",
    ])

    # ── Section 5: Special Cases ──────────────────────────────────────────────
    add_heading(doc, "5. Special Cases", 1)

    add_heading(doc, "5.1  WALI Sanitation", 2)
    add_body(doc, (
        "Sanitation costs and health benefits are not derived from a direct SSP output difference. "
        "Instead, per-capita factors are applied to the modelled population transitioning between "
        "sanitation tiers, differentiated by rural and urban setting. The factors below are in "
        "PPP 2019 USD per capita per year. Costs are negative (government investment); "
        "the health benefit is positive (avoided disease burden)."
    ))
    add_inline_table(doc,
        ["Sanitation tier", "Rural (USD/capita)", "Urban (USD/capita)"],
        [
            ["Unimproved",            "−6.5",   "−6.5"],
            ["Improved",              "−68.1",  "−34.1"],
            ["Safely managed",        "−102.1", "−66.2"],
            ["Health benefit (any tier)", "+200", "+200"],
        ],
        [6.0, 4.0, 4.0],
    )

    add_heading(doc, "5.2  Indoor Air Pollution", 2)
    add_body(doc, (
        "Indoor air pollution benefits monetise the health gains from households transitioning "
        "away from solid biomass cooking fuels (to gas or electricity). The method is based on "
        "the WHO BAR-HAP 2021 model, scaled to Uganda GDP per capita in PPP 2019 USD. "
        "The formula tracks the share of the population leaving each fuel type per year:"
    ))
    add_formula_box(doc, [
        "value (USD) = Σ population_affected(t) × net_benefit_per_capita",
        "",
        "  population_affected(t)  = transition_fraction(t) × total_population(t)",
        "  transition_fraction(t)  = max(frac_fuel(t) − frac_fuel(t+1), 0)",
        "  net_benefit_per_capita  = benefit_per_capita + cost_per_capita",
    ])
    add_body(doc, (
        "Cost and benefit factors (PPP 2019 USD per person, averaged across four policy "
        "instruments: stove subsidy, fuel subsidy, financing, and behavioural campaign):"
    ))
    add_inline_table(doc,
        ["Fuel transition", "Cost per capita (USD)", "Benefit per capita (USD)", "Net (USD)"],
        [
            ["Biomass → Gas",         "−6",  "+37", "+31"],
            ["Biomass → Electricity", "−19", "+57", "+38"],
        ],
        [5.0, 4.0, 4.5, 3.0],
    )

    # ── Appendix A: Capital Costs ──────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Appendix A — Capital Costs Reference", 1)
    add_body(doc, (
        "All capital cost line items (Tableau: 'Capital cost' category, broken down by Sector). "
        "Sorted by Tableau Sector; within each sector, investment costs appear before system costs."
    ))

    # Sort: sector order, then investment costs first, then display_name
    present_sectors = [s for s in SECTOR_ORDER if s in costs["tableau_category"].values]
    extra_sectors   = [s for s in costs["tableau_category"].unique() if s not in SECTOR_ORDER]
    sector_order_map = {s: i for i, s in enumerate(present_sectors + sorted(extra_sectors))}
    costs["_sector_ord"] = costs["tableau_category"].map(sector_order_map).fillna(999)
    costs["_src_ord"]    = costs["source"].apply(lambda x: 0 if "Investment" in str(x) else 1)
    costs_sorted = costs.sort_values(["_sector_ord", "_src_ord", "display_name"]).drop(
        columns=["_sector_ord", "_src_ord"]
    )

    col_defs_cost = [
        ("tableau_category", "Tableau Sector",  4.0),
        ("display_name",     "Cost Item",        5.5),
        ("source",           "Source",           4.5),
        ("factor_readable",  "Factor",           2.5),
        ("formula_type",     "Formula Type",     3.5),
    ]
    add_appendix_table(doc, costs_sorted, col_defs_cost, group_col="tableau_category")

    # ── Appendix B: Co-benefits ────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Appendix B — Co-benefits Reference", 1)
    add_body(doc, (
        "All co-benefit line items (all Tableau benefit categories except Capital cost). "
        "Sorted by Tableau Category; within each category, investment-linked items appear first."
    ))

    present_cats = [c for c in BENEFIT_ORDER if c in benefits["tableau_category"].values]
    extra_cats   = [c for c in benefits["tableau_category"].unique() if c not in BENEFIT_ORDER]
    cat_order_map = {c: i for i, c in enumerate(present_cats + sorted(extra_cats))}
    benefits["_cat_ord"] = benefits["tableau_category"].map(cat_order_map).fillna(999)
    benefits["_src_ord"] = benefits["source"].apply(lambda x: 0 if "Investment" in str(x) else 1)
    benefits_sorted = benefits.sort_values(["_cat_ord", "_src_ord", "display_name"]).drop(
        columns=["_cat_ord", "_src_ord"]
    )

    col_defs_ben = [
        ("tableau_category", "Tableau Category", 4.0),
        ("display_name",     "Benefit Item",      5.5),
        ("source",           "Source",            4.5),
        ("factor_readable",  "Factor",            2.5),
        ("formula_type",     "Formula Type",      3.5),
    ]
    add_appendix_table(doc, benefits_sorted, col_defs_ben, group_col="tableau_category")

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    build()
