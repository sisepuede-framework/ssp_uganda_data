"""
Render one of the docs/*.md review documents to a Word file.

Why
---
The change proposal and the RDM reference go to the wider team, who review in Word.
Pasting markdown into Word loses the tables, so we generate a real .docx instead.
`pandoc` is not installed on this machine; `python-docx` is, and the documents only
use a small, predictable subset of markdown — so a focused converter is enough.

Supported: # .. #### headings, paragraphs, - bullets, 1. numbered lists, > quotes,
pipe tables, **bold**, *italic*, `code`, --- rules, and screenshot placeholders.

Screenshots
-----------
A line like `> **[SCREENSHOT — Option 1]**` is replaced by `docs/images/option-1.png`
(or .jpg) when that file exists; otherwise the placeholder text is kept, so the
document can circulate before the designs are ready.

Usage
-----
    python backend/scripts/md_to_docx.py docs/change_proposal.md
    python backend/scripts/md_to_docx.py docs/change_proposal.md -o /somewhere/else.docx
"""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parents[2]
IMAGE_DIR = REPO_ROOT / "docs" / "images"

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
SCREENSHOT = re.compile(r"\[SCREENSHOT\s*[—-]\s*(.+?)\]")


def _slug(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")


def add_runs(paragraph, text: str) -> None:
    """Write text into a paragraph, honouring **bold**, *italic* and `code`."""
    text = text.replace("\\|", "|")
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def is_separator_row(line: str) -> bool:
    return bool(re.fullmatch(r"\|[\s\-:|]+\|", line))


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip("|").split("|")]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(width):
            text = row[j] if j < len(row) else ""
            cell = cells[j]
            cell.text = ""
            # <br> is used inside cells in the source documents.
            for k, chunk in enumerate(text.split("<br>")):
                para = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
                add_runs(para, chunk.strip())
                if i == 0:
                    for run in para.runs:
                        run.bold = True


def add_screenshot(doc: Document, label: str) -> None:
    """Insert docs/images/<slug>.png when it exists, else keep the placeholder."""
    slug = _slug(label)
    for ext in (".png", ".jpg", ".jpeg"):
        candidate = IMAGE_DIR / f"{slug}{ext}"
        if candidate.exists():
            doc.add_picture(str(candidate), width=Inches(6.2))
            caption = doc.add_paragraph(label)
            caption.style = "Caption"
            return
    para = doc.add_paragraph()
    run = para.add_run(f"[ screenshot to be inserted: {label} → docs/images/{slug}.png ]")
    run.italic = True


def convert(md_path: Path, out_path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    lines = md_path.read_text().splitlines()
    table_buffer: list[list[str]] = []
    i = 0

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, table_buffer)
            doc.add_paragraph()
            table_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # Tables accumulate until a non-table line ends them.
        if is_table_row(stripped):
            if not is_separator_row(stripped):
                table_buffer.append(split_row(stripped))
            i += 1
            continue
        flush_table()

        if not stripped or stripped in {"---", "***", "___"}:
            i += 1
            continue

        if stripped.startswith("<!--"):
            i += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            para = doc.add_heading("", level=level)
            add_runs(para, heading.group(2))
            i += 1
            continue

        if stripped.startswith(">"):
            body = stripped.lstrip("> ").strip()
            shot = SCREENSHOT.search(body)
            if shot:
                add_screenshot(doc, shot.group(1))
            else:
                para = doc.add_paragraph(style="Intense Quote")
                add_runs(para, body)
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            add_runs(para, bullet.group(1))
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            para = doc.add_paragraph(style="List Number")
            add_runs(para, numbered.group(1))
            i += 1
            continue

        if stripped.startswith("```"):          # skip fenced code fences, keep contents
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                para = doc.add_paragraph()
                run = para.add_run(lines[i])
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                i += 1
            i += 1
            continue

        # Plain paragraph: join the wrapped source lines that follow.
        block = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("#", "-", ">", "|", "*", "```"))
                    or re.match(r"^\d+\.\s", nxt)):
                break
            block.append(nxt)
            i += 1
        para = doc.add_paragraph()
        add_runs(para, " ".join(block))

    flush_table()
    doc.save(out_path)
    print(f"Wrote {out_path}")


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("markdown", type=Path, help="path to the .md file")
    ap.add_argument("-o", "--out", type=Path, default=None, help="output .docx path")
    args = ap.parse_args()

    md_path = args.markdown if args.markdown.is_absolute() else REPO_ROOT / args.markdown
    out_path = args.out or md_path.with_suffix(".docx")
    convert(md_path, out_path)


if __name__ == "__main__":
    main()
